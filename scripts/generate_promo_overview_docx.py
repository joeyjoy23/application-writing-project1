"""生成项目宣传用 Word 概要（基本情况、流程、亮点）。"""

from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from utils.export_word import (  # noqa: E402
    FONT_HEITI,
    FONT_SONG,
    _add_body,
    _add_center_title,
    _add_stage_heading,
    _add_subtitle,
    _set_page_margins,
    _set_run_font,
)


def _add_bullet(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    if bold_prefix and text.startswith(bold_prefix):
        r0 = p.add_run(bold_prefix)
        _set_run_font(r0, FONT_SONG, 12, bold=True)
        r1 = p.add_run(text[len(bold_prefix) :])
        _set_run_font(r1, FONT_SONG, 12)
    else:
        r = p.add_run("• " + text)
        _set_run_font(r, FONT_SONG, 12)


def _add_numbered(doc: Document, n: int, title: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    r0 = p.add_run(f"{n}. {title}")
    _set_run_font(r0, FONT_SONG, 12, bold=True)
    r1 = p.add_run(f"　{body}")
    _set_run_font(r1, FONT_SONG, 12)


def build_promo_document() -> Document:
    doc = Document()
    _set_page_margins(doc)

    _add_center_title(doc, "高考英语应用文 AI 分析系统", 22)
    _add_subtitle(
        doc,
        f"项目基本情况 · 工作流程 · 核心亮点\n（宣传素材整理稿 · {datetime.now().strftime('%Y年%m月%d日')}）",
    )

    _add_stage_heading(doc, "一、项目基本情况")
    _add_body(
        doc,
        "本产品是一款面向高考英语「应用文」写作的 AI 备课助手，采用固定四阶段工作流（Workflow），"
        "而非开放式聊天机器人。教师粘贴一道真题后，系统自动生成从审题分析、PEEL 建模、"
        "多版范文、语言支架到课堂教案的成套教研材料，支持导出 Word、历史存档与断点续跑。",
    )
    _add_body(doc, "定位：服务高中英语教师与教研组，用于个人备课、集体教研、课堂示范与讲评整理。")
    _add_body(doc, "形态：基于 Python + Streamlit 的 Web 应用，可本地运行或部署至 Streamlit Cloud。")
    _add_body(doc, "模型：支持 DeepSeek、智谱 GLM、OpenAI、Gemini、阿里百炼、小米 MiMo 等 OpenAI 兼容 API。")

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("产品类型", "多阶段 AI 备课工作流（非聊天）"),
        ("核心用户", "高中英语教师、教研员"),
        ("输入", "一道高考英语应用文真题（含题干与要求）"),
        ("输出", "审题蓝图 + PEEL + 三版范文 + 句型词汇 + 教学指南"),
        ("典型耗时", "完整四阶段约 8–15 分钟（视模型与网络而定）"),
    ]
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    _set_run_font(run, FONT_SONG, 11)

    _add_stage_heading(doc, "二、核心工作流程（四阶段）")
    _add_body(
        doc,
        "各阶段 Prompt 独立存放在 prompts/ 目录，由 workflow.py 动态加载，教研可改模板而无需改代码。"
        "Stage 2 与 Stage 3 支持并行执行以缩短等待时间。",
    )

    stages = [
        (
            "Stage 1　审题与结构分析",
            "输出 STRUCTURED_JSON（写作蓝图：体裁、要点、段落规划等）+ 教师可读总结。"
            "含动笔自查五问、交际任务三元分析、能力维度匹配、构思维度与要点结构规划。",
        ),
        (
            "Stage 2　PEEL 与多版范文",
            "基于 Stage 1 JSON 生成 PEEL 写作策略卡（核心/支撑要点便签）、"
            "要点覆盖表、基础版 + 情感共鸣型 + 逻辑思辨型两份高分范文、三版对比分析与升级解析。"
            "范文统一按 105–125 词撰写。",
        ),
        (
            "Stage 3　功能句型与话题词汇",
            "功能句型包（3 张语用功能表 + 本题适用说明与改句对照）+ "
            "话题词汇锦囊（按语义场分必备/进阶/亮点三级）。",
        ),
        (
            "Stage 4　教学指南与易错预警",
            "结合 Stage 1–3 与所选学生水平（基础/中等/进阶），生成学情适配教学路径、"
            "典型错误预警、课后练习题（含迁移题与可选用对比小练）。",
        ),
    ]
    for i, (title, body) in enumerate(stages, 1):
        _add_numbered(doc, i, title, body)

    flow_p = doc.add_paragraph()
    flow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = flow_p.add_run("原题  →  Stage 1  →  Stage 2 ∥ Stage 3  →  Stage 4  →  导出 Word / 历史存档")
    _set_run_font(r, FONT_SONG, 11)
    r.font.color.rgb = RGBColor(0x6B, 0x5B, 0x73)

    _add_stage_heading(doc, "三、项目亮点（宣传可提炼）")
    highlights = [
        (
            "教研级流水线，非聊天拼贴",
            "Stage 1 产出结构化 JSON 写作蓝图，后续阶段严格据此生成；要点覆盖可校验，"
            "避免「每次回答不同、易漏 mandatory 要点」。",
        ),
        (
            "范文有层次、有风格、有对比",
            "PEEL 便签区分核心/支撑要点；基础版 + 情感共鸣型 + 逻辑思辨型双高分路径；"
            "附三版对比表与高分升级点解析，便于课堂讲评。",
        ),
        (
            "语言支架与教案一体",
            "句型包按语用功能分表，配「本题宜用/不宜用」与改句示范；"
            "词汇按语义场分必备/进阶/亮点，例句贴合本题情境。",
        ),
        (
            "学情自适应教学指南",
            "Stage 4 按基础/中等/进阶生成差异化活动链、易错预警与分层练习，"
            "引用 Stage 2/3 具体句段与词块，杜绝空泛建议。",
        ),
        (
            "为真实课堂与运维设计",
            "Prompt 外置可迭代；多模型可切换；Stage 2/3 并行；可选 LLM 结果缓存降本；"
            "同题同模型历史自动合并；每完成一 Stage 即保存，失败可保留已完成部分。",
        ),
        (
            "可留存、可续用、可导出",
            "一键导出排版 Word（黑体标题 + 宋体正文）；云端可接 Neon 数据库持久化历史；"
            "历史记录支持载入续跑、收藏与管理员查看。",
        ),
    ]
    for i, (title, body) in enumerate(highlights, 1):
        _add_numbered(doc, i, title, body)

    _add_stage_heading(doc, "四、适用场景与目标受众")
    for item in [
        "课前个人备课：快速获得完整备课包，减少从零构思时间",
        "教研组集体备课：统一审题口径与范文标杆，讨论更有抓手",
        "课堂示范与写作讲评：三版范文对比、PEEL 便签可直接投影",
        "培优与分层教学：双风格高分范文 + 学情分级教学指南",
        "新教师上手应用文：动笔自查五问、易错预警降低常见误区",
    ]:
        _add_bullet(doc, item)

    _add_stage_heading(doc, "五、宣传海报可用金句（备选）")
    for q in [
        "一道真题，四套教研成果——审题、范文、句型、教案一次到位。",
        "不是聊天，是流水线：JSON 蓝图驱动，要点不漏、风格可讲。",
        "两份高分范文，两种思维路径：情感共鸣 × 逻辑思辨。",
        "PEEL 便签 + 三版对比，让应用文讲评有据可依。",
        "教师贴题，AI 备课；导出 Word，走进课堂。",
    ]:
        _add_bullet(doc, f"「{q}」")

    _add_stage_heading(doc, "六、技术与部署概要")
    _add_body(
        doc,
        "技术栈：Python · Streamlit · OpenAI Compatible API · python-docx · "
        "SQLite（本地）/ PostgreSQL Neon（云端）。",
    )
    _add_body(
        doc,
        "仓库：GitHub 托管，Streamlit Cloud 一键部署；Secrets 配置 API Key、DATABASE_URL、ADMIN_PASSWORD。",
    )
    _add_body(doc, "质量保障：pytest 覆盖解析、工作流序列化、进度推断与 Markdown 排版等模块。")

    _add_subtitle(doc, "—— 文档结束 · 可根据各节拆分为宣传长图、海报与介绍页 ——")

    return doc


def main() -> None:
    out_dir = ROOT / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "项目宣传-基本情况流程与亮点.docx"
    doc = build_promo_document()
    doc.save(out_path)
    print(f"已生成：{out_path}")


if __name__ == "__main__":
    main()
