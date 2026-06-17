"""生成面向学生使用者的功能介绍与亮点宣传 Word 文档。"""

from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.generate_promo_overview_docx import (  # noqa: E402
    _add_bullet,
    _add_numbered,
)
from utils.export_word import (  # noqa: E402
    FONT_SONG,
    _add_body,
    _add_center_title,
    _add_stage_heading,
    _add_subtitle,
    _set_page_margins,
    _set_run_font,
)


def build_student_promo_document() -> Document:
    doc = Document()
    _set_page_margins(doc)

    _add_center_title(doc, "高考英语应用文 AI 学习助手", 22)
    _add_subtitle(
        doc,
        f"学生功能介绍 · 学习亮点 · 使用建议\n（面向学生使用者 · {datetime.now().strftime('%Y年%m月%d日')}）",
    )

    _add_stage_heading(doc, "一、这是什么？")
    _add_body(
        doc,
        "这是一款围绕高考英语「应用文」写作设计的 AI 学习工具。老师输入一道真题后，"
        "系统会按固定步骤生成一整套学习材料——从「这道题考什么」到「怎么写、用什么词、"
        "怎么练」，不是随便聊天的机器人，而是一条清晰的学习流水线。",
    )
    _add_body(
        doc,
        "作为学生，你通常通过老师分享的 Word 备课包、课堂投影或只读分享链接来使用这些内容；"
        "你也可以在老师开放的网页版中，对照自己的习作与系统给出的范文、句型与易错点进行自学。",
    )

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("适合谁", "高一至高三，备考高考英语应用文的同学"),
        ("解决什么", "审题不清、结构松散、词汇贫乏、不知道高分范文长什么样"),
        ("你能拿到", "审题导图、PEEL 策略、三版范文、句型词汇包、易错提醒与练习"),
        ("范文标准", "每篇范文 105–125 英文词（高于部分真题字数要求，利于积累表达）"),
        ("怎么用", "跟老师课堂 + 课后对照范文与句型自主消化"),
    ]
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    _set_run_font(run, FONT_SONG, 11)

    _add_stage_heading(doc, "二、四步学习流水线（你会看到什么）")
    _add_body(
        doc,
        "老师点一次「完整流程」后，系统依次（部分步骤并行）生成下面四块内容。"
        "你可以按 Stage 1→4 的顺序自学，也可以只重点看范文和句型词汇。",
    )

    stages = [
        (
            "第一步　吃透题目（审题）",
            "帮你弄清：写什么体裁、给谁写、要覆盖哪些要点、用什么时态和人称。"
            "相当于动笔前的「导航图」，避免漏要点、跑题。",
        ),
        (
            "第二步　学会结构 + 看范文（PEEL 与三版范文）",
            "PEEL 策略卡：每个要点怎么展开（观点—解释—例证—收束）。"
            "三版范文：基础版（稳拿分）+ 情感共鸣型高分范文 + 逻辑思辨型高分范文，"
            "并附三版对比与「高分升级点」解析，方便你模仿不同风格。",
        ),
        (
            "第三步　积累语言（句型与词汇）",
            "功能句型包：建议信、演讲稿等常用功能句，标注「本题宜用 / 不宜用」并给改句示范。"
            "话题词汇锦囊：按「必备 / 进阶 / 亮点」分级，例句贴合本题情境，可直接背记迁移。",
        ),
        (
            "第四步　针对你水平的学法（教学指南与练习）",
            "老师可选择基础 / 中等 / 进阶学生水平，生成对应的学法建议、"
            "易错预警（你最容易错在哪）和课后练习题，帮你从「看懂」到「写对」。",
        ),
    ]
    for i, (title, body) in enumerate(stages, 1):
        _add_numbered(doc, i, title, body)

    flow_p = doc.add_paragraph()
    flow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = flow_p.add_run("真题  →  审题  →  范文与结构  →  句型词汇  →  易错与练习")
    _set_run_font(r, FONT_SONG, 11)
    r.font.color.rgb = RGBColor(0x6B, 0x5B, 0x73)

    _add_stage_heading(doc, "三、六大学习亮点（为什么值得用）")
    highlights = [
        (
            "审题有清单，不再漏要点",
            "系统把题目拆成可检查的要点表，并提示体裁、语域与段落安排。"
            "写之前对照一遍，减少「写了半页发现漏了一个要点」的失误。",
        ),
        (
            "范文不只一篇，而是「基础 + 双高分路径」",
            "同一道题给出三种完整范文：稳扎稳打的基础版，以及情感向、逻辑向两种高分思路。"
            "你可以先模仿基础版，再挑战高分版的表达与结构。",
        ),
        (
            "PEEL 便签，结构不再散",
            "每个要点对应一张 PEEL 便签，告诉你这一段先写什么、后写什么。"
            "特别适合习惯「想到哪写到哪」的同学建立段落意识。",
        ),
        (
            "句型词汇「能直接用」",
            "不是泛泛的词表，而是按本题情境筛选的句型与词汇，"
            "并说明哪些表达适合这道题、哪些容易用错。",
        ),
        (
            "易错预警，提前避坑",
            "针对本题常见错误（时态、人称、要点遗漏、中式英语等）给出提醒，"
            "相当于把老师口头的「大家注意」写成了可反复看的清单。",
        ),
        (
            "分层学习，难度可匹配",
            "基础 / 中等 / 进阶不同学法路径，避免「全班同一套要求」带来的挫败或吃不饱。",
        ),
    ]
    for i, (title, body) in enumerate(highlights, 1):
        _add_numbered(doc, i, title, body)

    _add_stage_heading(doc, "四、学生怎么用（三步上手）")
    _add_numbered(
        doc,
        1,
        "课堂跟老师",
        "老师投影或讲解系统生成的审题表、PEEL 便签与范文对比；你重点记「要点清单」和「高分升级点」。",
    )
    _add_numbered(
        doc,
        2,
        "课后对照消化",
        "打开老师分享的 Word 或链接，先默读三版范文，圈出你想模仿的句型；"
        "再到「句型词汇」板块摘抄 5–10 个表达，写进自己的积累本。",
    )
    _add_numbered(
        doc,
        3,
        "动笔 + 自评",
        "用同一真题限时写作，写完后对照要点表自评是否覆盖全部 mandatory 要点，"
        "再对照易错预警改一版；有余力则挑战情感型或逻辑型范文中的复杂句式。",
    )

    _add_stage_heading(doc, "五、推荐学习节奏（以一道题为例）")
    for item in [
        "第 1 天：精读审题总结 + 划出三个必写要点",
        "第 2 天：精读基础版范文，仿写一段开头和结尾",
        "第 3 天：学习功能句型包，替换自己习作中的简单句",
        "第 4 天：背诵「进阶 / 亮点」词汇各 5 个，并造句",
        "第 5 天：限时完整写作 + 对照易错预警修改",
        "第 6 天（可选）：精读高分范文 A 或 B，尝试升级一个段落",
    ]:
        _add_bullet(doc, item)

    _add_stage_heading(doc, "六、适合你的学习场景")
    for item in [
        "考前冲刺：快速积累一道真题的范文与高频表达",
        "平时周练：每道应用文作业都走一遍「审题 → 范文 → 句型」",
        "培优拓展：对比两种高分范文，训练思维深度与语言档次",
        "订正复盘：用易错预警对照自己的旧作文，标出同类错误",
        "小组学习：分工精读三版范文，小组汇报「高分升级点」",
    ]:
        _add_bullet(doc, item)

    _add_stage_heading(doc, "七、宣传海报可用金句（学生向）")
    for q in [
        "一道真题，看清要点、看懂范文、带走句型。",
        "不只背范文，更学会 PEEL：观点、解释、例子、收束一步到位。",
        "基础版稳分，双高分范文带你看见「情感」与「逻辑」两条路。",
        "必备、进阶、亮点词汇分级背，写考场不再词穷。",
        "易错预警先知道，考场上少丢冤枉分。",
    ]:
        _add_bullet(doc, f"「{q}」")

    _add_stage_heading(doc, "八、温馨提示")
    _add_body(
        doc,
        "本系统生成的是学习参考与语言支架，不能替代你自己的思考与动笔。"
        "建议「先理解再模仿，先仿写再创新」；范文用于学习结构与表达，"
        "考试时须结合题目独立成文，切忌照搬整段。",
    )
    _add_body(
        doc,
        "若老师提供网页分享链接，一般为只读预览，便于在手机端浏览；"
        "完整学习与打印建议以老师导出的 Word 备课包为准。",
    )

    _add_subtitle(doc, "—— 祝备考顺利 · 笔下有思路，卷面有亮点 ——")

    return doc


def main() -> None:
    out_dir = ROOT / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "学生宣传-功能介绍与亮点.docx"
    doc = build_student_promo_document()
    doc.save(out_path)
    print(f"已生成：{out_path}")


if __name__ == "__main__":
    main()
