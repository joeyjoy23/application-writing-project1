"""将应用文备课四阶段结果导出为排版友好的 Word 文档。"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Iterator

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from utils.parsers import sanitize_llm_html_breaks, strip_reader_self_check

FONT_HEITI = "黑体"
FONT_SONG = "宋体"
FONT_HEITI_EN = "SimHei"
FONT_SONG_EN = "SimSun"
FONT_ENGLISH = "Times New Roman"
SIZE_ENGLISH_PT = 12  # 小四号
COLOR_ENGLISH = RGBColor(0, 0, 0)  # 黑色

_INLINE_CODE_RE = re.compile(r"`([^`]+)`")
_ENGLISH_RUN_RE = re.compile(
    r"[A-Za-z](?:[A-Za-z0-9\s\-',.;:!?\"()\[\]/%&@#+*=<>]*[A-Za-z0-9]|[A-Za-z]{2,})",
)

STAGE_TITLES = {
    1: "Stage 1  审题与结构分析",
    2: "Stage 2  PEEL 与多版范文",
    3: "Stage 3  功能句型与话题词汇",
    4: "Stage 4  教学指南与易错预警",
}


def _set_run_font(
    run,
    font_name: str,
    size_pt: int,
    *,
    bold: bool = False,
    color: RGBColor | None = None,
) -> None:
    run.font.name = FONT_HEITI_EN if font_name == FONT_HEITI else FONT_SONG_EN
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), run.font.name)
    r_fonts.set(qn("w:hAnsi"), run.font.name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _set_run_english(run, *, bold: bool = False) -> None:
    run.font.name = FONT_ENGLISH
    run.font.size = Pt(SIZE_ENGLISH_PT)
    run.font.bold = bold
    run.font.color.rgb = COLOR_ENGLISH
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT_ENGLISH)
    r_fonts.set(qn("w:hAnsi"), FONT_ENGLISH)
    r_fonts.set(qn("w:eastAsia"), FONT_ENGLISH)


def _is_english_segment(text: str) -> bool:
    s = text.strip()
    if len(s) < 2:
        return False
    latin = sum(1 for c in s if c.isascii() and c.isalpha())
    cjk = sum(1 for c in s if "\u4e00" <= c <= "\u9fff")
    if latin >= 2 and latin >= cjk:
        return True
    return latin >= 3 and cjk == 0


def _split_english_runs(text: str) -> list[tuple[bool, str]]:
    segments: list[tuple[bool, str]] = []
    last = 0
    for m in _ENGLISH_RUN_RE.finditer(text):
        if m.start() > last:
            segments.append((False, text[last : m.start()]))
        segments.append((True, m.group(0)))
        last = m.end()
    if last < len(text):
        segments.append((False, text[last:]))
    return segments


def _split_text_segments(text: str) -> list[tuple[bool, str]]:
    segments: list[tuple[bool, str]] = []
    pos = 0
    for m in _INLINE_CODE_RE.finditer(text):
        if m.start() > pos:
            segments.extend(_split_english_runs(text[pos : m.start()]))
        segments.append((True, m.group(1)))
        pos = m.end()
    if pos < len(text):
        segments.extend(_split_english_runs(text[pos:]))
    return segments


def _add_runs_to_paragraph(p, text: str, *, body_size: int = 12, bold: bool = False) -> None:
    plain = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    plain = re.sub(r"\*([^*\n]+?)\*", r"\1", plain)
    for is_english, segment in _split_text_segments(plain):
        if not segment:
            continue
        run = p.add_run(segment)
        if is_english or _is_english_segment(segment):
            _set_run_english(run, bold=bold)
        else:
            _set_run_font(run, FONT_SONG, body_size, bold=bold)


def _set_page_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)


def _strip_inline_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*([^*\n]+?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return bool(s) and s.startswith("|") and s.endswith("|")


def _is_separator_row(line: str) -> bool:
    stripped = line.strip().strip("|")
    if not stripped:
        return False
    cells = [c.strip() for c in stripped.split("|")]
    return all(re.fullmatch(r":?-{3,}:?", c) for c in cells if c)


def _parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _parse_md_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if _is_separator_row(line):
            continue
        row = _parse_table_row(line)
        if row:
            rows.append([_strip_inline_md(c) for c in row])
    return rows


def _heading_level(line: str) -> int | None:
    m = re.match(r"^(#{1,4})\s+(.+)$", line.strip())
    if m:
        return len(m.group(1))
    return None


def _heading_text(line: str) -> str:
    m = re.match(r"^#{1,4}\s+(.+)$", line.strip())
    return _strip_inline_md(m.group(1)) if m else _strip_inline_md(line)


def _is_list_item(line: str) -> bool:
    return bool(re.match(r"^(\s*[-*•]|\s*\d+\.)\s+", line.strip()))


def _list_text(line: str) -> str:
    return _strip_inline_md(re.sub(r"^(\s*[-*•]|\s*\d+\.)\s+", "", line.strip()))


def _iter_markdown_blocks(text: str) -> Iterator[tuple[str, Any]]:
    lines = text.split("\n")
    i = 0
    buf: list[str] = []

    def flush_para() -> Iterator[tuple[str, Any]]:
        nonlocal buf
        if buf:
            joined = "\n".join(buf).strip()
            buf = []
            if joined:
                yield ("paragraph", joined)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            yield from flush_para()
            i += 1
            continue

        level = _heading_level(line)
        if level is not None:
            yield from flush_para()
            yield (f"heading{min(level, 4)}", _heading_text(line))
            i += 1
            continue

        if _is_table_row(line):
            yield from flush_para()
            table_lines: list[str] = []
            while i < len(lines) and _is_table_row(lines[i]):
                table_lines.append(lines[i])
                i += 1
            yield ("table", _parse_md_table(table_lines))
            continue

        if _is_list_item(line):
            yield from flush_para()
            items: list[str] = []
            while i < len(lines) and _is_list_item(lines[i]):
                items.append(_list_text(lines[i]))
                i += 1
            yield ("list", items)
            continue

        if stripped.startswith("```"):
            yield from flush_para()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            yield ("code", "\n".join(code_lines).strip())
            continue

        buf.append(line)
        i += 1

    yield from flush_para()


def _add_center_title(doc: Document, text: str, size_pt: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    _set_run_font(run, FONT_HEITI, size_pt, bold=True)


def _add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    _set_run_font(run, FONT_SONG, 11, color=RGBColor(0x55, 0x55, 0x55))


def _add_stage_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    _set_run_font(run, FONT_HEITI, 16, bold=True)


def _add_heading(doc: Document, text: str, level: int) -> None:
    sizes = {1: 15, 2: 14, 3: 13, 4: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_run_font(run, FONT_HEITI, sizes.get(level, 12), bold=True)


def _add_body(doc: Document, text: str, *, first_line_indent: bool = False) -> None:
    clean = _strip_inline_md(text)
    if not clean.strip():
        return
    for part in clean.split("\n"):
        part = part.strip()
        if not part:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Pt(24) if first_line_indent else Pt(0)
        _add_runs_to_paragraph(p, part)


def _add_code_block(doc: Document, text: str) -> None:
    if not text.strip():
        return
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.15)
        run = p.add_run(line.rstrip())
        _set_run_english(run)


def _add_list(doc: Document, items: list[str]) -> None:
    for item in items:
        if not item.strip():
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        _add_runs_to_paragraph(p, item)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            bold = r_idx == 0
            _add_runs_to_paragraph(cell.paragraphs[0], cell_text, body_size=11, bold=bold)
    doc.add_paragraph()


def _write_markdown(doc: Document, text: str, *, indent_paragraphs: bool = False) -> None:
    if not text or not text.strip():
        _add_body(doc, "（暂无内容）")
        return
    for block_type, data in _iter_markdown_blocks(text):
        if block_type.startswith("heading"):
            level = int(block_type.replace("heading", "") or "2")
            _add_heading(doc, data, min(level, 4))
        elif block_type == "paragraph":
            _add_body(doc, data, first_line_indent=indent_paragraphs)
        elif block_type == "code":
            _add_code_block(doc, data)
        elif block_type == "list":
            _add_list(doc, data)
        elif block_type == "table":
            _add_table(doc, data)


def export_workflow_to_word(
    *,
    question: str,
    stage1_summary: str | None = None,
    stage2_raw: str | None = None,
    stage3_raw: str | None = None,
    stage4_raw: str | None = None,
) -> bytes:
    """生成 .docx 字节流。"""
    doc = Document()
    _set_page_margins(doc)

    _add_center_title(doc, "高考英语应用文备课分析报告", 22)
    _add_subtitle(
        doc,
        f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}",
    )

    _add_stage_heading(doc, "题目原文")
    _add_body(doc, question.strip() or "（未填写）")

    if stage1_summary:
        _add_stage_heading(doc, STAGE_TITLES[1])
        _write_markdown(doc, strip_reader_self_check(stage1_summary))

    if stage2_raw:
        _add_stage_heading(doc, STAGE_TITLES[2])
        _write_markdown(
            doc, strip_reader_self_check(stage2_raw), indent_paragraphs=True
        )

    if stage3_raw:
        _add_stage_heading(doc, STAGE_TITLES[3])
        _write_markdown(doc, sanitize_llm_html_breaks(stage3_raw))

    if stage4_raw:
        _add_stage_heading(doc, STAGE_TITLES[4])
        _write_markdown(doc, stage4_raw)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
